"""
transyt.py — Download a YouTube transcript to a formatted Word document.

Usage:
    python transyt.py <youtube_url> [-l LANGUAGE] [-o OUTPUT_DIR]
    python transyt.py  # prompts for URL interactively
"""

import argparse
import html
import json
import os
import re
import subprocess
import sys
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from youtube_transcript_api import (
    NoTranscriptFound,
    TranscriptsDisabled,
    VideoUnavailable,
    YouTubeRequestFailed,
    YouTubeTranscriptApi,
)

TARGET_CHUNK_CHARS = 350
DEFAULT_LANGUAGES = ["en", "en-US", "en-GB"]
BODY_FONT = "Calibri"
MONO_FONT = "Courier New"
BODY_SIZE_PT = 11
TIMESTAMP_SIZE_PT = 9


def extract_video_id(url: str) -> str:
    pattern = r"(?:youtube\.com/(?:watch\?v=|shorts/|embed/)|youtu\.be/)([a-zA-Z0-9_-]{11})"
    match = re.search(pattern, url)
    if not match:
        raise ValueError(
            f"Could not extract a video ID from: {url}\n"
            "Expected formats: https://youtube.com/watch?v=..., https://youtu.be/..., "
            "https://youtube.com/shorts/..., https://youtube.com/embed/..."
        )
    return match.group(1)


def fetch_metadata(video_id: str) -> dict:
    url = f"https://youtube.com/watch?v={video_id}"
    fallback = {"title": video_id, "upload_date": None, "url": url, "duration": None}
    try:
        result = subprocess.run(
            ["yt-dlp", "--dump-json", "--no-playlist", "--no-warnings", "--", url],
            capture_output=True,
            text=True,
            timeout=30,
        )
        data = json.loads(result.stdout)
        raw_date = data.get("upload_date")
        formatted_date = None
        if raw_date:
            try:
                formatted_date = datetime.strptime(raw_date, "%Y%m%d").strftime("%B %d, %Y")
            except ValueError:
                formatted_date = raw_date
        return {
            "title": data.get("title", video_id),
            "upload_date": formatted_date,
            "url": data.get("webpage_url", url),
            "duration": data.get("duration"),
        }
    except FileNotFoundError:
        print("Warning: yt-dlp not found. Video title will use the video ID.")
    except subprocess.TimeoutExpired:
        print("Warning: Metadata fetch timed out. Using video ID as title.")
    except (json.JSONDecodeError, KeyError, Exception):
        print("Warning: Could not parse video metadata. Using video ID as title.")
    return fallback


def fetch_transcript(video_id: str, languages: list) -> tuple:
    # v1.x API is instance-based
    api = YouTubeTranscriptApi()
    transcript_list = api.list(video_id)

    # Prefer manually created captions; fall back to auto-generated; then any language
    try:
        transcript = transcript_list.find_manually_created_transcript(languages)
    except NoTranscriptFound:
        try:
            transcript = transcript_list.find_generated_transcript(languages)
        except NoTranscriptFound:
            transcript = next(iter(transcript_list))
            print(
                f"Warning: Requested language not available. "
                f"Using '{transcript.language_code}' instead."
            )

    fetched = transcript.fetch()
    seg_list = [{"text": s.text, "start": s.start, "duration": s.duration} for s in fetched]
    return seg_list, transcript.language_code


def group_segments(segments: list, target: int = TARGET_CHUNK_CHARS) -> list:
    paragraphs = []
    chunk_text = ""
    chunk_start = 0.0

    for seg in segments:
        text = html.unescape(seg["text"]).strip()
        if not text:
            continue
        if chunk_text and (len(chunk_text) + 1 + len(text)) > target:
            paragraphs.append({"start": chunk_start, "text": chunk_text})
            chunk_text = text
            chunk_start = seg["start"]
        else:
            chunk_text = (chunk_text + " " + text).strip() if chunk_text else text
            if not paragraphs and not chunk_text:
                chunk_start = seg["start"]
            elif chunk_text == text:
                chunk_start = seg["start"]

    if chunk_text:
        paragraphs.append({"start": chunk_start, "text": chunk_text})

    return paragraphs


def format_timestamp(seconds: float) -> str:
    total = int(seconds)
    h = total // 3600
    m = (total % 3600) // 60
    s = total % 60
    if h:
        return f"{h}:{m:02d}:{s:02d}"
    return f"{m}:{s:02d}"


def sanitize_filename(title: str) -> str:
    sanitized = re.sub(r'[<>:"/\\|?*]', "", title)
    sanitized = re.sub(r"\s+", "_", sanitized).strip("._")
    sanitized = sanitized[:200]
    return sanitized or "transcript"


def _add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_meta_line(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run_label = p.add_run(label)
    run_label.bold = True
    run_label.font.name = BODY_FONT
    run_label.font.size = Pt(9)
    run_value = p.add_run(value)
    run_value.font.name = BODY_FONT
    run_value.font.size = Pt(9)


def build_document(metadata: dict, paragraphs: list, lang_code: str) -> Document:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Title
    title_p = doc.add_paragraph(metadata["title"], style="Title")
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacer

    # Metadata block
    _add_meta_line(doc, "URL: ", metadata["url"])
    if metadata["upload_date"]:
        _add_meta_line(doc, "Date: ", metadata["upload_date"])
    _add_meta_line(doc, "Language: ", lang_code)

    _add_horizontal_rule(doc)

    heading = doc.add_heading("Transcript", level=1)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)

    # Transcript paragraphs
    for para in paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(14)

        run_ts = p.add_run(f"[{format_timestamp(para['start'])}]  ")
        run_ts.font.name = MONO_FONT
        run_ts.font.size = Pt(TIMESTAMP_SIZE_PT)
        run_ts.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run_ts.bold = True

        run_body = p.add_run(para["text"])
        run_body.font.name = BODY_FONT
        run_body.font.size = Pt(BODY_SIZE_PT)

    return doc


def save_document(doc: Document, title: str, output_dir: str = ".") -> str:
    base = sanitize_filename(title)
    path = os.path.join(output_dir, base + ".docx")
    counter = 2
    while os.path.exists(path):
        path = os.path.join(output_dir, f"{base}_{counter}.docx")
        counter += 1
    doc.save(path)
    return os.path.abspath(path)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Download a YouTube transcript to a Word document",
        epilog="Example: python transyt.py https://youtu.be/dQw4w9WgXcQ",
    )
    parser.add_argument("url", nargs="?", help="YouTube video URL")
    parser.add_argument(
        "-l", "--language", default="en", help="Preferred transcript language (default: en)"
    )
    parser.add_argument(
        "-o", "--output-dir", default=".", help="Output directory (default: current directory)"
    )
    args = parser.parse_args()

    url = args.url
    if not url:
        url = input("Enter YouTube video URL: ").strip()
    if not url:
        print("Error: No URL provided.")
        sys.exit(1)

    languages = [args.language, "en", "en-US", "en-GB"]
    languages = list(dict.fromkeys(languages))  # deduplicate while preserving order

    try:
        video_id = extract_video_id(url)
    except ValueError as e:
        print(f"Error: {e}")
        sys.exit(1)

    print("Fetching metadata...")
    metadata = fetch_metadata(video_id)

    print("Fetching transcript...")
    try:
        segments, lang_code = fetch_transcript(video_id, languages)
    except TranscriptsDisabled:
        print("Error: Transcripts are disabled for this video.")
        sys.exit(1)
    except VideoUnavailable:
        print("Error: Video is unavailable (private or deleted).")
        sys.exit(1)
    except NoTranscriptFound:
        print("Error: No transcript found for this video in any language.")
        sys.exit(1)
    except YouTubeRequestFailed as e:
        print(f"Error: YouTube request failed. Check your internet connection.\n{e}")
        sys.exit(1)
    except Exception as e:
        print(f"Error fetching transcript: {e}")
        sys.exit(1)

    print("Building document...")
    grouped = group_segments(segments)
    doc = build_document(metadata, grouped, lang_code)

    output_path = save_document(doc, metadata["title"], args.output_dir)
    print(f"Saved transcript to: {output_path}")


if __name__ == "__main__":
    main()
